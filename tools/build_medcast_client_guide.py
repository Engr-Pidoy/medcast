from __future__ import annotations

import json
import os
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = Path(
    os.getenv(
        "MEDCAST_GUIDE_OUTPUT",
        ROOT / "docs" / "MEDCAST_Client_User_and_Technical_Guide.docx",
    )
)
BENCHMARK = ROOT / "storage" / "app" / "medcast" / "benchmark_result.json"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17365D"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F3F7FB"
GOLD = "C59B3A"
PALE_GOLD = "FFF7E0"
GREEN = "DDEBF7"
GRAY = "666666"
LIGHT_GRAY = "F2F2F2"
WHITE = "FFFFFF"
RED = "C00000"
CONTENT_DXA = 9360


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = tr_pr.find(qn("w:cantSplit"))
        if cant_split is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[min(idx, len(widths) - 1)] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[min(idx, len(widths) - 1)]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("222222")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Title", 28, NAVY, 0, 12),
        ("Subtitle", 13, GRAY, 0, 12),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    styles["Title"].font.bold = True
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.bold = True

    for name in ("List Bullet", "List Number"):
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st.paragraph_format.left_indent = Inches(0.375)
        st.paragraph_format.first_line_indent = Inches(-0.188)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25

    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Callout"]
    callout.font.name = "Calibri"
    callout.font.size = Pt(10.5)
    callout.font.color.rgb = RGBColor.from_string(NAVY)
    callout.paragraph_format.left_indent = Inches(0.15)
    callout.paragraph_format.right_indent = Inches(0.15)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.line_spacing = 1.15

    if "Small Note" not in styles:
        small = styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        small = styles["Small Note"]
    small.font.name = "Calibri"
    small.font.size = Pt(9)
    small.font.color.rgb = RGBColor.from_string(GRAY)
    small.paragraph_format.space_after = Pt(4)
    small.paragraph_format.line_spacing = 1.1

    if "Table Text" not in styles:
        table_text = styles.add_style("Table Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        table_text = styles["Table Text"]
    table_text.font.name = "Calibri"
    table_text.font.size = Pt(9.5)
    table_text.paragraph_format.space_before = Pt(0)
    table_text.paragraph_format.space_after = Pt(2)
    table_text.paragraph_format.line_spacing = 1.05

    doc.core_properties.title = "MEDCAST Client User and Technical Guide"
    doc.core_properties.subject = "Client guide for the MEDCAST hospital admissions forecasting web application"
    doc.core_properties.author = "MEDCAST Project Team"
    doc.core_properties.keywords = "MEDCAST, forecasting, admissions, capacity risk, model evaluation"
    doc.core_properties.comments = "Prepared from the local MEDCAST application and its current evaluation output."


def configure_header_footer(section, first_page: bool = False) -> None:
    section.different_first_page_header_footer = first_page
    header = section.header
    p = header.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("MEDCAST  |  Client User & Technical Guide")
    r.font.name = "Calibri"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(GRAY)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("MEDCAST  •  Local documentation  •  Page ")
    r.font.name = "Calibri"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(GRAY)
    add_page_number(fp)
    if first_page:
        first_header = section.first_page_header.paragraphs[0]
        first_header.clear()
        first_footer = section.first_page_footer.paragraphs[0]
        first_footer.clear()


def add_border(paragraph, color=BLUE, size=12, space=1) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(size))
    top.set(qn("w:space"), str(space))
    top.set(qn("w:color"), color)
    p_bdr.append(top)


def add_heading(doc, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def add_body(doc, text: str, bold_lead: str | None = None, style: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items: list[str]):
    paragraphs = []
    for item in items:
        paragraphs.append(doc.add_paragraph(item, style="List Bullet"))
    return paragraphs


def add_numbers(doc, items: list[str]) -> None:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level.append(level_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)

    for item in items:
        paragraph = doc.add_paragraph(item, style="List Number")
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_node])
        paragraph._p.get_or_add_pPr().append(num_pr)


def add_callout(doc, title: str, text: str, tone: str = "blue") -> None:
    fill = PALE_BLUE if tone == "blue" else PALE_GOLD
    accent = BLUE if tone == "blue" else GOLD
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.style = doc.styles["Callout"]
    r = p.add_run(f"{title}  ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(accent)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int], font_size: float = 9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.style = doc.styles["Table Text"]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(NAVY)
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.style = doc.styles["Table Text"]
            p.add_run(str(value))
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def page_break(doc) -> None:
    doc.add_page_break()


def fmt(value, digits=2, suffix="") -> str:
    if value is None:
        return "—"
    if isinstance(value, (int, float)):
        return f"{value:.{digits}f}{suffix}"
    return str(value)


def load_snapshot() -> dict:
    if not BENCHMARK.exists():
        return {}
    with BENCHMARK.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def build() -> Path:
    snapshot = load_snapshot()
    doc = Document()
    configure_document(doc)
    configure_header_footer(doc.sections[0], first_page=True)

    # Cover
    for _ in range(3):
        doc.add_paragraph()
    accent = doc.add_paragraph()
    accent.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_border(accent, color=GOLD, size=22, space=4)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("MEDCAST")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Client User & Technical Guide")
    descriptor = doc.add_paragraph()
    descriptor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = descriptor.add_run("Hospital Admissions Forecasting, Model Evaluation,\nand Capacity-Risk Decision Support")
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    descriptor.paragraph_format.space_after = Pt(28)
    box = doc.add_table(rows=3, cols=2)
    box.style = "Table Grid"
    set_table_geometry(box, [2700, 6660])
    cover_rows = [
        ("Audience", "Hospital administrators, researchers, data encoders, reviewers, and project clients"),
        ("Application", "MEDCAST research web application, version 1.2.0"),
        ("Guide status", "Local working documentation; values may change when the dataset is updated and forecasts are rerun"),
    ]
    for i, (label, value) in enumerate(cover_rows):
        set_cell_shading(box.cell(i, 0), LIGHT_BLUE)
        p = box.cell(i, 0).paragraphs[0]
        p.style = doc.styles["Table Text"]
        p.add_run(label).bold = True
        p2 = box.cell(i, 1).paragraphs[0]
        p2.style = doc.styles["Table Text"]
        p2.add_run(value)
    doc.add_paragraph()
    note = doc.add_paragraph(style="Small Note")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run(f"Prepared {date.today().strftime('%B %Y')}  •  NDH scope  •  For research and planning support")

    page_break(doc)

    add_heading(doc, "How to use this guide", 1)
    add_body(doc, "This document is the primary client reference for what MEDCAST does, what each page displays, how values are computed, and how results should—and should not—be interpreted. Use the Quick Answers section first, then consult the page guide, metric reference, and FAQ for detail.")
    add_callout(doc, "Important boundary", "MEDCAST supports operational planning. It does not diagnose patients, prescribe treatment, replace hospital policy, or automatically authorize staffing and bed decisions.", "gold")

    add_heading(doc, "Guide map", 2)
    add_table(
        doc,
        ["Section", "What it answers"],
        [
            ["1. Quick Answers", "The most likely questions from the client or panel"],
            ["2. System Overview", "Purpose, users, data flow, and scope"],
            ["3. Core Assumptions", "80/20 split, 120-bed basis, thresholds, and horizons"],
            ["4. Page-by-Page Guide", "What every page and control means"],
            ["5–7. Models and Evaluation", "Selection logic, metrics, intervals, and sensitivity analyses"],
            ["8–9. Capacity Risk and Decision Support", "Risk calculations and safe operational interpretation"],
            ["10–13. Operations and Reference", "Data rules, workflow, limitations, FAQs, and glossary"],
        ],
        [2300, 7060],
    )

    add_heading(doc, "Document conventions", 2)
    add_bullets(doc, [
        "A dash (—) means the metric is not defined or there was not enough applicable data for that calculation.",
        "Percentages are rounded for display; calculations use the underlying values.",
        "“Current snapshot” means the latest locally generated output bundled with the app at the time this guide was produced.",
        "Forecasts, best-model rankings, thresholds, and risk results can change after data entry, import, or rerun.",
    ])

    page_break(doc)

    add_heading(doc, "1. Quick Answers for the Client", 1)
    quick_rows = [
        ["What is MEDCAST?", "A research web application that forecasts daily hospital admissions, evaluates several forecasting models, summarizes bed occupancy, and converts forecast results into planning-oriented risk indicators."],
        ["What hospital is represented?", "The current implementation is scoped to NDH and uses the NDH hospital code internally."],
        ["What data split is used?", "Chronological 80% initial training and 20% rolling-origin evaluation. Within the newest 20%, each daily origin adds newly revealed observations to an expanding training window before the next refit."],
        ["Why 120 beds?", "The occupancy calculation uses a mean operational capacity of 120 beds: 100 regular/base beds plus overflow or additional beds used when demand exceeds the regular count."],
        ["Can occupied beds exceed 100?", "Yes. The source discussion explains that overflow beds may be added. That is why the operational assumption is 120 rather than limiting every day to 100."],
        ["Is daily admissions capacity also 120?", "No. Capacity Risk uses an admission-handling capacity per day derived from historical admission percentiles or a custom input. Bed capacity and daily admission capacity are different measures."],
        ["Why is the best model not always SARIMA?", "Model choice is evidence-based. MEDCAST compares candidate models across expanding-window rolling origins and selects the best model per horizon. SARIMA remains a benchmark and powers the separate monthly research outlook."],
        ["What do 1, 7, and 30 days mean?", "They are forecast horizons: tomorrow, the next week, and the next 30 days. Accuracy and uncertainty commonly differ by horizon."],
        ["Can the result be used as an official directive?", "No. It is decision support. Hospital leaders must validate current census, staffing, policy, outbreaks, scheduled services, and other context before acting."],
    ]
    add_table(doc, ["Question", "Answer"], quick_rows, [2600, 6760], font_size=9.0)

    add_heading(doc, "Current local snapshot", 2)
    if snapshot:
        best_by_horizon = {}
        for item in snapshot.get("benchmarks", []):
            if item.get("is_best_for_horizon"):
                best_by_horizon[item.get("horizon_days")] = item
        best_text = ", ".join(
            f"{h}-day: {best_by_horizon[h].get('model_name', '—')}"
            for h in sorted(best_by_horizon)
        ) or "Not available"
        add_table(doc, ["Item", "Current value"], [
            ["Dataset version", snapshot.get("dataset_version", "—")],
            ["Coverage", f"{snapshot.get('dataset_coverage_start', '—')} to {snapshot.get('dataset_coverage_end', '—')}"],
            ["Records", f"{snapshot.get('dataset_records', '—')} daily records"],
            ["Evaluation design", f"{snapshot.get('training_records', '—')} initial training (80%) / {snapshot.get('testing_records', '—')} rolling evaluation days (20%)"],
            ["Best model by horizon", best_text],
            ["Selected SARIMA order", snapshot.get("selected_sarima_order", "—")],
        ], [2700, 6660])
    else:
        add_body(doc, "No local benchmark snapshot was available when this guide was generated.")
    add_callout(doc, "Do not memorize the snapshot", "Treat these values as traceability information, not permanent specifications. Rerunning the pipeline after a data update may produce a different dataset version, threshold, best model, or score.")

    page_break(doc)

    add_heading(doc, "2. System Overview", 1)
    add_heading(doc, "2.1 Purpose", 2)
    add_body(doc, "MEDCAST organizes historical admissions data, creates short-horizon forecasts, evaluates competing statistical models, monitors occupancy context, and presents planning signals in one authenticated web application. Its main benefit is traceability: the user can see the forecast, the uncertainty interval, the historical basis, the model metrics, and the assumptions behind a risk label.")

    add_heading(doc, "2.2 Intended users", 2)
    add_bullets(doc, [
        "Hospital administrator or operations lead — reviews demand, occupancy, capacity pressure, and preparedness guidance.",
        "Researcher or thesis reviewer — inspects rolling-origin comparisons, prediction-interval behavior, and sensitivity analyses.",
        "Data encoder — enters or imports daily counts and checks that totals and dates are correct.",
        "System administrator — controls access, verifies the deployment, backs up data, and triggers model reruns when appropriate.",
    ])

    add_heading(doc, "2.3 End-to-end data flow", 2)
    add_numbers(doc, [
        "Daily data are encoded manually or imported from a CSV/Excel file.",
        "MEDCAST stores the historical observations and validates the date/count fields.",
        "The forecasting pipeline converts the series to daily frequency and prepares missing dates if needed.",
        "The oldest 80% establishes the initial training window; the newest 20% defines the rolling-origin evaluation period.",
        "At each daily origin, the training history expands with newly revealed observations, each model is refitted, and 1-, 7-, and 30-day-ahead forecasts are evaluated where targets are available.",
        "Naive, Seasonal Naive, SARIMA, Prophet, and Holt-Winters are compared using aggregate errors across all eligible origins for each horizon.",
        "The app stores forecasts, prediction intervals, evaluation metrics, thresholds, and model metadata.",
        "Dashboard, Forecasting, Performance, Capacity Risk, and Decision Support pages present different views of the same analytical workflow.",
    ])

    page_break(doc)
    add_heading(doc, "2.4 Technology overview", 2)
    add_table(doc, ["Layer", "Role"], [
        ["Laravel / PHP", "Application routes, authentication, data management, and server-side orchestration"],
        ["Livewire + Blade", "Interactive screens and reusable interface components"],
        ["Python forecasting pipeline", "Time-series preparation, model fitting, forecasting, benchmarking, and interval metrics"],
        ["Database + generated artifacts", "Historical records, forecast runs, benchmark outputs, figures, and traceability metadata"],
    ], [2500, 6860])

    add_heading(doc, "3. Core Assumptions and Definitions", 1)
    add_heading(doc, "3.1 Chronological 80/20 rolling-origin evaluation", 2)
    add_body(doc, "The dataset is split by time, not randomly. The first 80% establishes the initial model-training window, while the most recent 20% defines the evaluation period. Starting at that boundary, MEDCAST advances the forecast origin one day at a time. After each actual observation becomes available, it is added to the expanding training window before the models are refitted for the next origin.")
    add_callout(doc, "Why chronological?", "Randomly mixing future and past observations can leak information and make time-series accuracy look better than it really is. Chronological splitting preserves order.")
    add_callout(doc, "Eligible origins", "A 1-day horizon uses every evaluation origin. A 7-day or 30-day horizon uses fewer origins because enough future actual observations must remain to score that lead time. The app reports the origin count for every horizon.", "gold")

    add_heading(doc, "3.2 Capacity terms that must not be mixed", 2)
    add_table(doc, ["Term", "Meaning", "Current basis"], [
        ["Regular/base bed capacity", "Beds normally available without overflow arrangements", "100 beds"],
        ["Mean operational bed capacity", "Planning denominator that includes typical overflow/additional beds", "120 beds"],
        ["Occupied beds", "Actual beds occupied by patients on a date", "May exceed 100; can also exceed 120 during severe pressure"],
        ["Bed occupancy rate", "Occupied beds ÷ 120 × 100", "Operational occupancy context"],
        ["Admission-handling capacity", "Number of new admissions per day the selected scenario assumes can be handled", "Historical P50, P66, P90, or a custom value"],
        ["Capacity-pressure ratio", "Average forecast admissions per day ÷ admission-handling capacity per day", "Used on Capacity Risk"],
    ], [2200, 4400, 2760], font_size=8.7)
    add_callout(doc, "Example", "If 126 beds are occupied, operational occupancy is 126 ÷ 120 × 100 = 105%. This does not mean there were 126 new admissions that day. Occupancy is a census; admissions are a daily flow.", "gold")

    add_heading(doc, "3.3 Demand levels", 2)
    add_body(doc, "Demand classes are derived from the training distribution so the thresholds reflect the historical NDH admissions pattern:")
    add_table(doc, ["Class", "Rule", "Interpretation"], [
        ["Low", "Forecast ≤ 33rd percentile (P33)", "Lower portion of historical daily demand"],
        ["Moderate", "P33 < forecast ≤ 66th percentile (P66)", "Middle portion of historical daily demand"],
        ["High", "Forecast > P66", "Upper portion of historical daily demand; review preparedness"],
    ], [1800, 3200, 4360])
    if snapshot:
        thresholds = snapshot.get("thresholds", {}).get("meta", {})
        add_body(doc, f"Current local thresholds: P33 = {thresholds.get('p33', '—')} and P66 = {thresholds.get('p66', '—')} admissions/day. These values are regenerated from the active dataset.", style="Small Note")

    add_heading(doc, "3.4 Forecast horizons", 2)
    add_table(doc, ["Horizon", "Typical use", "Caution"], [
        ["1 day", "Immediate coordination for tomorrow", "Very short window; one unusual day can dominate the metric"],
        ["7 days", "Weekly planning, monitoring, and model selection", "Uncertainty accumulates across days"],
        ["30 days", "Medium-term preparedness and scenario review", "Use ranges and patterns; do not treat each point as exact"],
    ], [1500, 3900, 3960])

    page_break(doc)

    add_heading(doc, "4. Page-by-Page User Guide", 1)
    add_heading(doc, "4.1 Dashboard", 2)
    add_body(doc, "The Dashboard is the operational summary. It brings together the latest historical status, the active forecast run, and the indicators most likely to require attention.")
    add_table(doc, ["Dashboard element", "How to read it"], [
        ["KPI cards", "Latest admissions, occupied beds/occupancy, forecast totals or averages, and trend summaries depending on available data"],
        ["7-day forecast", "Point forecast plus uncertainty bands for the upcoming week"],
        ["Tomorrow demand", "Next-day forecast classified as Low, Moderate, or High from the percentile thresholds"],
        ["Admissions-type breakdown", "Regular, emergency, and other counts where source data provide those categories"],
        ["Recent admissions", "Newest encoded/imported records for quick validation"],
        ["Trend/occupancy visuals", "Context for whether demand is rising and whether beds are under pressure"],
    ], [2800, 6560])
    add_callout(doc, "Client check", "Confirm the latest date shown is expected. A convincing forecast based on stale data is still operationally stale.")

    add_heading(doc, "4.2 Encode Daily Data", 2)
    add_body(doc, "Use this page to add one day of data. The total admissions value is computed from the admissions categories.")
    add_table(doc, ["Field", "Meaning / rule"], [
        ["Date", "Observation date; avoid duplicate dates"],
        ["Regular admissions", "Non-emergency admissions counted for the day"],
        ["Emergency admissions", "Admissions through emergency pathways"],
        ["Other admissions", "Admissions that do not fit the first two categories"],
        ["Discharges", "Patients discharged that day"],
        ["Occupied beds", "Actual occupied-bed census; may exceed the regular 100-bed base"],
        ["Notes", "Optional operational context such as outbreaks, events, or data corrections"],
        ["Rerun forecast", "If selected, refreshes models/results after saving; this can take time"],
    ], [2400, 6960])

    add_heading(doc, "4.3 Historical Data", 2)
    add_body(doc, "This page supports record review and bulk data import. Monthly aggregation helps users inspect longer patterns. An import merges rows by date: new dates are added and matching dates are updated. The merged records are stored in the persistent database and normally trigger a new forecasting run, so the uploaded file must be validated first.")
    add_bullets(doc, [
        "Back up the current dataset before a large merge or correction.",
        "Confirm the date column and admissions column are mapped correctly.",
        "Check for duplicate dates, gaps, negative values, text in numeric fields, and unexpected totals.",
        "After import, confirm dataset version, coverage dates, record count, and latest date on Model Performance or Dashboard.",
    ])

    add_heading(doc, "4.4 Trends", 2)
    add_body(doc, "Trends summarizes average admissions by weekday and month, overall mean/minimum/maximum/variation, and a fitted slope. “Increasing” or “decreasing” describes the direction of the historical series; it does not prove a cause or guarantee that the trend will continue.")

    add_heading(doc, "4.5 Forecasting", 2)
    add_body(doc, "The Forecasting page is the detailed forecast view. Select a 1-, 7-, or 30-day horizon, then review the model, daily point forecasts, and prediction intervals. Use the model-comparison area to see how alternatives performed.")
    add_table(doc, ["Element", "Meaning"], [
        ["Selected/active model", "Model chosen for that horizon from rolling-origin performance"],
        ["Point forecast", "Single best estimate for a future day"],
        ["80% prediction interval", "Narrower plausible range with lower target coverage"],
        ["95% prediction interval", "Wider plausible range with higher target coverage"],
        ["Model comparison", "Candidate errors and interval characteristics for the selected horizon"],
        ["Monthly SARIMA outlook", "Separate research-oriented 12-month average-daily-admissions projection from the selected SARIMA specification"],
    ], [2700, 6660])
    add_callout(doc, "Best practice", "For operations, report a range and scenario—not only the point forecast. The longer the horizon, the more important the interval becomes.")

    add_heading(doc, "4.6 Model Performance", 2)
    add_body(doc, "This page provides the audit trail behind the model choice. It contains rolling-origin counts and dates, basic errors, prediction-interval quality, high-demand classification metrics, error robustness, SARIMA order selection, sensitivity analyses, and dataset version/coverage information.")
    add_bullets(doc, [
        "Use MAE/MASE to compare average forecast error; use RMSE when large errors deserve extra attention.",
        "Review coverage together with width. Very wide intervals can achieve high coverage without being operationally useful.",
        "Review sensitivity and missed-event rate when failing to flag high demand is costly.",
        "Review specificity and false-alert rate when unnecessary escalation is costly.",
        "Use the dataset version and split metadata to make results reproducible in the paper or presentation.",
    ])

    add_heading(doc, "4.7 Capacity Risk", 2)
    add_body(doc, "Capacity Risk is a scenario page. It does not change the hospital’s real capacity. It asks: given a forecast horizon, an assumed daily admission-handling capacity, and a planning penalty preference, how much forecast pressure and overload risk should the user prepare for?")

    add_heading(doc, "4.8 Decision Support", 2)
    add_body(doc, "Decision Support combines forecast demand levels, occupancy context, emergency share, and rule-based recommendations. It helps users translate analytical results into a review checklist. Recommendations are general planning suggestions and require human validation.")

    add_heading(doc, "4.9 About", 2)
    add_body(doc, "The About page summarizes the application version, purpose, technical context, and research disclaimer. Use it when explaining the system’s scope to new users or reviewers.")

    add_heading(doc, "5. Forecasting Models and Selection", 1)
    add_table(doc, ["Model", "What it represents", "Why it is included"], [
        ["Naive", "Tomorrow resembles the latest observation", "Simple baseline; complex models should justify their added complexity"],
        ["Seasonal Naive", "Tomorrow resembles the same weekday in the prior weekly cycle", "Baseline for weekly hospital patterns"],
        ["SARIMA", "Autoregressive, differenced, moving-average model with weekly seasonality", "Classical time-series benchmark and monthly outlook source"],
        ["Prophet", "Trend/seasonality model designed for time-indexed data", "Alternative that can model smooth trend and recurring patterns"],
        ["Holt-Winters", "Exponential smoothing with level, trend, and seasonal components", "Responsive classical alternative for recurring seasonal demand"],
    ], [1900, 4400, 3060], font_size=8.8)

    add_heading(doc, "5.1 How the best model is chosen", 2)
    add_body(doc, "Candidate models are evaluated through an expanding-window rolling-origin procedure inside the chronological 20% evaluation period. For each eligible origin, the model is refitted using only information available by that date and the exact 1-, 7-, or 30-day-ahead prediction is scored. The app aggregates the errors across origins and identifies a best model for each horizon, using scaled error (MASE) as the primary comparison and MAE/RMSE as supporting evidence. The primary operational model is tied to the 7-day comparison because weekly planning is the central use case.")
    add_callout(doc, "Important", "A model name is not a permanent winner. New data, a different split, or a changed outlier pattern can alter the ranking. This is expected behavior, not an error.")

    add_heading(doc, "5.2 SARIMA order selection", 2)
    add_body(doc, "The pipeline evaluates candidate SARIMA specifications and records information criteria such as AIC and BIC. The evaluation specification is selected using only the initial 80% training set, then its parameters are refitted at each rolling origin. A separate full-history selection supports the current live forecast. This separation prevents future evaluation observations from choosing the benchmark specification.")

    add_heading(doc, "5.3 Missing-day preparation", 2)
    add_body(doc, "The Python pipeline regularizes the observations to a daily series. If a date is missing, the series is completed and missing values are interpolated, with backward/forward filling as safeguards. This supports model execution, but repeated gaps should be corrected at the source because imputation is not equivalent to observed hospital data.")

    add_heading(doc, "5.4 Current best-model snapshot", 2)
    if snapshot:
        rows = []
        for item in snapshot.get("benchmarks", []):
            if item.get("is_best_for_horizon"):
                rows.append([
                    f"{item.get('horizon_days')} days",
                    item.get("model_name", "—"),
                    fmt(item.get("mae")),
                    fmt(item.get("rmse")),
                    fmt(item.get("mase"), 3),
                    fmt(item.get("coverage_80"), 1, "%"),
                ])
        add_table(doc, ["Horizon", "Best model", "MAE", "RMSE", "MASE", "80% coverage"], rows, [1200, 1900, 1450, 1450, 1450, 1910], font_size=8.6)

    add_heading(doc, "6. Metric Reference", 1)
    add_heading(doc, "6.1 Forecast error metrics", 2)
    add_table(doc, ["Metric", "Plain-language meaning", "Interpretation"], [
        ["MAE", "Average absolute difference between forecast and actual admissions", "Lower is better; expressed in admissions"],
        ["RMSE", "Square-root average of squared errors", "Lower is better; penalizes large misses more heavily"],
        ["MASE", "MAE scaled against a simple naive benchmark", "Lower is better; below 1 generally means better than the scale benchmark"],
        ["High-demand MAE", "MAE calculated only on actual high-demand days", "Lower is better; exposes performance when pressure matters most"],
    ], [1550, 4700, 3110], font_size=8.8)

    add_heading(doc, "6.2 Prediction-interval metrics", 2)
    add_table(doc, ["Metric", "Plain-language meaning", "Interpretation"], [
        ["Coverage by horizon", "Percentage of actual values that fell inside the stated interval", "Compare with the 80% or 95% target; context and sample size matter"],
        ["Average interval width", "Average upper bound minus lower bound", "Narrower is more precise if coverage remains adequate"],
        ["Relative interval width", "Average width divided by mean actual demand", "Normalizes width so horizons/models are easier to compare"],
    ], [2100, 4350, 2910], font_size=8.8)
    add_callout(doc, "Coverage trade-off", "Coverage alone is not enough. A 100% coverage result may come from an interval so wide that it provides little planning value. Judge coverage and width together.", "gold")

    add_heading(doc, "6.3 High-demand event metrics", 2)
    add_table(doc, ["Metric", "Formula", "Question answered"], [
        ["Sensitivity (Recall)", "TP ÷ (TP + FN)", "Of actual high-demand days, how many did the model flag?"],
        ["Specificity", "TN ÷ (TN + FP)", "Of actual non-high-demand days, how many were correctly left unflagged?"],
        ["Precision", "TP ÷ (TP + FP)", "Of days flagged high demand, how many truly were high?"],
        ["F1-score", "2 × precision × sensitivity ÷ (precision + sensitivity)", "How balanced are precision and sensitivity?"],
        ["False-alert rate", "FP ÷ (FP + TN)", "How often were normal days incorrectly flagged?"],
        ["Missed-event rate", "FN ÷ (FN + TP)", "How often were actual high-demand days missed?"],
    ], [1900, 2900, 4560], font_size=8.6)
    add_body(doc, "TP = true positive, TN = true negative, FP = false positive, and FN = false negative. A metric may be undefined when its denominator is zero—for example, if the evaluated window contains no actual high-demand day.", style="Small Note")

    add_heading(doc, "6.4 Rolling-origin robustness", 2)
    add_body(doc, "MEDCAST orders the absolute errors produced by the daily rolling origins and computes a short rolling average over those errors. It summarizes the rolling MAE mean and variation. The robustness score is 100 × (1 − min(1, coefficient of variation of rolling MAE)). A higher score means the errors were more stable across origins; it does not automatically mean the errors were small.")

    page_break(doc)
    add_heading(doc, "7. Strong Evaluation and Sensitivity Analysis", 1)
    add_heading(doc, "7.1 Capacity sensitivity", 2)
    add_body(doc, "The evaluation tests multiple admission-capacity thresholds derived from the training data: P50 (median), P66, and P90. For each threshold it compares actual and predicted exceedance days and forecast overload. This shows whether conclusions change under a constrained, moderate, or expanded operating assumption.")

    add_heading(doc, "7.2 Threshold sensitivity", 2)
    add_body(doc, "High-demand event metrics are recalculated at P60, P66, and P75. A lower threshold flags more days and may improve sensitivity while increasing false alerts. A higher threshold focuses on more extreme days but can miss moderately elevated demand.")

    add_heading(doc, "7.3 Penalty sensitivity", 2)
    add_table(doc, ["Mode", "Priority", "Evaluation weighting"], [
        ["Balanced", "Equal concern for missed events and false alerts", "Missed event 1.0; false alert 1.0"],
        ["Overload-sensitive", "Missing a high-demand event is more costly", "Missed event 2.0; false alert 0.5"],
        ["Resource-conservative", "Unnecessary escalation is more costly", "Missed event 0.75; false alert 2.0"],
    ], [2300, 4200, 2860])
    add_body(doc, "These evaluation weights are related to—but not identical to—the Capacity Risk page multipliers described in Section 8.", style="Small Note")

    add_heading(doc, "7.4 Outlier sensitivity", 2)
    add_body(doc, "The app compares baseline MAE with MAE after capping extreme values at the training-series 99th percentile (winsorization). A large change suggests the reported error is sensitive to rare extreme days. This is diagnostic only; it does not silently remove the original observations from normal reporting.")

    add_heading(doc, "7.5 High-demand-day performance", 2)
    add_body(doc, "The high-demand threshold used for the principal event analysis is the training-series P66. The app separately reports error on actual high-demand days and the classification metrics. This prevents strong performance on routine days from hiding poor performance during operational peaks.")

    add_callout(doc, "How to present results", "Show at least one error metric, interval coverage plus width, high-demand sensitivity/missed-event rate, and a sensitivity result. Avoid declaring a model “accurate” from one number alone.")

    page_break(doc)
    add_heading(doc, "8. Capacity-Risk Scenarios", 1)
    add_heading(doc, "8.1\u00a0Controls", 2)
    add_table(doc, ["Control", "Options", "Effect"], [
        ["Forecast horizon", "1, 7, or 30 days", "Chooses the matching best model and future forecast window"],
        ["Capacity scenario", "Constrained (P50), Moderate (P66), Expanded (P90), or Custom", "Sets the assumed number of admissions that can be handled per day"],
        ["Penalty setting", "Balanced, Overload-sensitive, or Resource-conservative", "Adjusts how strongly the final risk score responds"],
    ], [2100, 3350, 3910])

    add_heading(doc, "8.2 Displayed results", 2)
    add_table(doc, ["Result", "Calculation / meaning"], [
        ["Selected forecasting model", "Best model stored for the chosen horizon and latest valid forecast run"],
        ["Forecasted admissions", "Average point forecast per day over the selected horizon"],
        ["Assumed scenario capacity", "Chosen daily admission-handling capacity—not the number of beds"],
        ["Capacity-pressure ratio", "Average forecast admissions ÷ scenario capacity; above 1.0 means average demand exceeds the assumption"],
        ["Expected overload", "Sum across days of max(forecast − capacity, 0)"],
        ["Probability of exceeding capacity", "Approximate chance of at least one exceedance during the horizon"],
        ["Capacity-risk level", "Low, Moderate, High, or Critical from the risk score"],
        ["Recommended preparedness", "Routine monitoring, standby, enhanced readiness, or emergency readiness"],
    ], [2750, 6610], font_size=8.8)

    add_heading(doc, "8.3 Probability and risk calculation", 2)
    add_body(doc, "Per-day exceedance probability is approximated from the point forecast and the 80% prediction interval under a normal-error assumption. Horizon-level probability combines the daily probabilities using an independence approximation. The risk score gives 65% weight to exceedance probability and 35% to capacity pressure, then applies the selected penalty multiplier.")
    add_table(doc, ["Penalty", "Risk multiplier"], [
        ["Balanced", "1.00"],
        ["Overload-sensitive", "1.15"],
        ["Resource-conservative", "0.85"],
    ], [4000, 5360])
    page_break(doc)
    add_table(doc, ["Risk score", "Level", "Preparedness interpretation"], [
        ["Below 25", "Low", "Routine monitoring"],
        ["25 to below 50", "Moderate", "Standby / review near-term capacity"],
        ["50 to below 75", "High", "Enhanced readiness and active coordination"],
        ["75 or higher", "Critical", "Emergency readiness and leadership review"],
    ], [2200, 2000, 5160])
    add_callout(doc, "Approximation warning", "Daily admission forecasts are related across time, so the independence assumption may overstate or understate the true horizon probability. Use the result as a scenario indicator, not a guaranteed probability.", "gold")

    add_heading(doc, "8.4 Recommended interpretation sequence", 2)
    add_numbers(doc, [
        "Check the selected horizon and model.",
        "Confirm the scenario capacity represents admissions per day, not beds.",
        "Read pressure ratio and expected overload before the risk label.",
        "Compare at least two capacity scenarios and penalty settings.",
        "Validate against current occupancy, staffing, scheduled services, emergency conditions, and hospital policy.",
        "Record the chosen scenario and rationale if the result is used in a report or meeting.",
    ])

    add_heading(doc, "9. Decision Support and Operational Use", 1)
    add_heading(doc, "9.1 What the page combines", 2)
    add_bullets(doc, [
        "Upcoming forecast days and their Low/Moderate/High demand classes",
        "Operational bed occupancy using the 120-bed mean capacity assumption",
        "Emergency-admissions share where categorized data are available",
        "Rule-based alerts and general preparedness recommendations",
    ])

    add_heading(doc, "9.2 Safe workflow for a planning meeting", 2)
    add_numbers(doc, [
        "Validate data freshness and the latest observation date.",
        "Review the point forecast together with the 80% and 95% intervals.",
        "Check whether high-demand flags agree with known schedules or events.",
        "Review current occupancy using 120 as the documented operational denominator.",
        "Open Capacity Risk and test constrained, moderate, expanded, and—if justified—custom scenarios.",
        "Document non-model factors such as staffing, supplies, outbreaks, closures, and referral patterns.",
        "Make and approve the operational decision through normal hospital governance.",
    ])

    add_heading(doc, "9.3 What MEDCAST does not know automatically", 2)
    add_table(doc, ["Factor", "Why human validation is required"], [
        ["Staff skill mix and absences", "A bed or admission count does not reveal whether the right personnel are available"],
        ["Elective schedules and closures", "Future known events may not yet appear in historical data"],
        ["Disease outbreaks or disasters", "Structural changes can invalidate historical relationships"],
        ["Supplies and equipment", "Capacity may be constrained by resources other than beds"],
        ["Patient acuity and length of stay", "Aggregate admissions do not capture all workload and bed-turnover differences"],
        ["Policy and referral changes", "A new process can change the data-generating pattern"],
    ], [3000, 6360], font_size=8.8)

    add_callout(doc, "Escalation principle", "A High or Critical label should trigger review and coordination. It should not trigger an irreversible action without current operational confirmation and authorized approval.")

    add_heading(doc, "10. Data Entry and Import Reference", 1)
    add_heading(doc, "10.1 Minimum data quality rules", 2)
    add_bullets(doc, [
        "Use one record per calendar date and maintain chronological coverage.",
        "Use non-negative numeric counts for admissions, discharges, and occupied beds.",
        "Make total admissions equal to regular + emergency + other when categories are supplied.",
        "Do not convert missing observations to zero unless zero is the verified true value.",
        "Record exceptional events in notes and preserve the source file used for import.",
        "Investigate abrupt jumps before assuming they are errors; they may represent genuine high-demand events.",
    ])

    add_heading(doc, "10.2 Import expectations", 2)
    add_body(doc, "CSV and Excel imports should contain a recognizable date field and an admissions field. The import logic accepts common header aliases, but a simple, explicit layout is safest. If only aggregate admissions are available, category-level charts may place the total in one category while emergency/other remain zero; this is a source-data limitation, not a forecast failure.")
    add_table(doc, ["Recommended column", "Example", "Notes"], [
        ["date", "2024-08-18", "ISO date format is preferred"],
        ["regular_admissions", "25", "Use if category detail is available"],
        ["emergency_admissions", "8", "Use if category detail is available"],
        ["other_admissions", "2", "Optional category"],
        ["total_admissions", "35", "Required if categories are not provided"],
        ["discharges", "30", "Optional but useful operational context"],
        ["occupied_beds", "118", "Actual occupied-bed census"],
        ["notes", "Festival weekend", "Optional context"],
    ], [2600, 2400, 4360], font_size=8.8)

    add_heading(doc, "10.3 Pre-import checklist", 2)
    add_numbers(doc, [
        "Copy the source file and keep an unchanged backup.",
        "Check row count, first date, last date, duplicate dates, and missing dates.",
        "Check minimum, maximum, mean, and impossible negative values.",
        "Confirm whether occupied beds include overflow beds.",
        "Confirm whether admissions are aggregate or categorized.",
        "Import, then compare MEDCAST record count and coverage with the source.",
    ])

    add_heading(doc, "10.4 When to rerun forecasting", 2)
    add_bullets(doc, [
        "After importing or updating the dataset",
        "After correcting material historical errors",
        "After adding a meaningful block of new daily observations",
        "Before a formal report when the displayed run may be stale",
        "After model-code or assumption changes that affect evaluation",
    ])

    add_heading(doc, "11. Administration, Traceability, and Troubleshooting", 1)
    add_heading(doc, "11.1 Traceability items to record", 2)
    add_table(doc, ["Item", "Why it matters"], [
        ["Dataset version", "Connects screenshots, tables, and conclusions to one exact input dataset"],
        ["Coverage period and record count", "Shows how much history was used and how current it is"],
        ["Training/testing split", "Documents the evaluation design"],
        ["Forecast run date/batch", "Identifies which result the user reviewed"],
        ["Selected model and horizon", "Prevents mixing 1-, 7-, and 30-day results"],
        ["Capacity and penalty scenario", "Makes Capacity Risk results reproducible"],
        ["Known data caveats", "Explains category gaps, overflow beds, or unusual events"],
    ], [2900, 6460])

    add_heading(doc, "11.2 Troubleshooting", 2)
    add_table(doc, ["Symptom", "Likely cause", "What to do"], [
        ["No forecast shown", "No valid forecast run or insufficient/invalid data", "Check historical records, rerun, and review server/pipeline logs"],
        ["Latest date is wrong", "Stale or incorrectly formatted import", "Check source dates and re-import only after backup"],
        ["Forecast changed after upload", "Models were retrained on a new dataset", "Expected; compare dataset version and evaluation metrics"],
        ["Metric shows —", "Denominator is zero or applicable events are absent", "Review the horizon/confusion matrix; do not convert undefined to 0%"],
        ["Occupancy exceeds 100%", "Occupied beds exceed the 120 mean operational capacity", "Confirm census and overflow use; treat as pressure, not an automatic data error"],
        ["Occupied beds exceed 100 beds", "Overflow/additional beds are represented", "Use the documented 120-bed operational assumption"],
        ["Capacity Risk seems inconsistent with beds", "Scenario capacity is admissions/day, not bed count", "Recheck units and the selected scenario"],
        ["Rerun takes time", "Five models are repeatedly refitted across daily rolling origins", "Avoid repeated clicks; wait for completion and check the latest batch"],
    ], [2500, 3000, 3860], font_size=8.2)

    add_heading(doc, "11.3 Access and deployment hygiene", 2)
    access_bullets = add_bullets(doc, [
        "Keep the application behind authentication and limit accounts to authorized users.",
        "Change demo/default credentials before client or production use.",
        "Do not upload patient-identifying information; MEDCAST is designed for aggregate daily counts.",
        "Back up the database and source dataset before major imports or upgrades.",
        "Test locally before any GitHub push or Render deployment, and deploy only after explicit approval.",
    ])
    for paragraph in access_bullets[:-1]:
        paragraph.paragraph_format.keep_with_next = True

    add_heading(doc, "12. Limitations and Responsible Use", 1)
    add_table(doc, ["Limitation", "Practical implication"], [
        ["Research prototype", "Requires local validation, governance, security review, and operational acceptance before real-world reliance"],
        ["Single-hospital scope", "Findings and thresholds are specific to the current NDH data and assumptions"],
        ["Aggregate daily admissions", "Does not fully represent acuity, ward-level demand, service line, or length of stay"],
        ["Category availability", "Imported aggregate data may not support meaningful regular/emergency/other comparisons"],
        ["One historical evaluation period", "Rolling-origin refitting is stronger than one fixed-origin forecast, but performance can still differ in another calendar period or under a structural change"],
        ["Prediction-interval approximations", "Some model intervals and the scaled Prophet 95% interval rely on simplifying assumptions"],
        ["Capacity probability approximation", "Normal-error and daily-independence assumptions may not hold exactly"],
        ["Historical dependence", "Outbreaks, policy changes, closures, or new referral patterns can reduce relevance of older data"],
        ["Capacity assumptions", "120 beds is a documented mean operational planning value, not a guarantee that all 120 are staffed or usable every day"],
        ["Synchronous rolling reruns", "Daily-origin refitting is computationally intensive; model generation may take several minutes and should not be repeatedly triggered"],
    ], [3000, 6360], font_size=8.5)

    add_heading(doc, "12.1 Required disclaimer", 2)
    add_callout(doc, "For reports and demonstrations", "MEDCAST forecasts and recommendations are analytical decision-support outputs. They must be interpreted with current hospital information and approved through normal clinical and administrative governance. The system does not provide medical diagnosis or treatment advice.", "gold")

    add_heading(doc, "12.2 Statements to avoid", 2)
    add_table(doc, ["Avoid saying", "Prefer saying"], [
        ["“The model guarantees 40 admissions.”", "“The point forecast is 40, with an uncertainty range that should guide scenario planning.”"],
        ["“SARIMA is always the best model.”", "“SARIMA is one benchmark; the app selects the best-performing model for each horizon.”"],
        ["“The hospital has exactly 120 permanent beds.”", "“The app uses 120 as mean operational capacity, reflecting a 100-bed base plus overflow.”"],
        ["“Critical risk means we must automatically add staff.”", "“Critical risk should trigger immediate validation and authorized preparedness review.”"],
        ["“100% coverage proves perfect forecasts.”", "“Coverage must be assessed together with interval width and the number of tested observations.”"],
    ], [3950, 5410], font_size=8.7)

    add_heading(doc, "13. Frequently Asked Questions", 1)
    faqs = [
        ("Why use 80% initial training and 20% rolling evaluation?", "It preserves a substantial recent evaluation period while keeping time order intact. Within the newest 20%, each origin only uses observations that would have been available at that point."),
        ("What does rolling-origin mean?", "The forecast origin advances one day at a time. Newly observed actual data are added to an expanding training window, the models are refitted, and the requested lead-time forecast is scored against its future actual value."),
        ("Why are there fewer 30-day than 1-day origins?", "A 30-day forecast can only be scored when 30 future actual observations remain. The current 156-day evaluation period therefore provides 156 one-day, 150 seven-day, and 127 thirty-day origins."),
        ("Why does the current best model differ from an earlier result?", "Changing the split or adding/correcting data changes the evidence available to each model. A new winner is a legitimate result when the same comparison rules are applied."),
        ("Why keep SARIMA if another model wins?", "SARIMA is central to the research comparison, provides interpretable seasonal structure, and supports the monthly outlook. Keeping it as a benchmark strengthens the evaluation."),
        ("What does an 80% prediction interval mean?", "Across repeated forecasts made under similar assumptions, the method aims for roughly 80% of actual outcomes to fall inside the interval. It does not mean one specific day has an absolute guarantee."),
        ("Why is the 95% interval wider?", "Higher target coverage requires a broader range. The extra width reflects greater caution, not lower quality by itself."),
        ("Is 100% interval coverage ideal?", "Not automatically. It may indicate conservative, very wide intervals. Compare coverage with average and relative width."),
        ("Why can sensitivity or precision be blank?", "The chosen horizon may contain no actual or predicted high-demand events, making the denominator zero. Blank/— is mathematically more honest than forcing 0%."),
        ("Why can occupancy be above 100%?", "The actual occupied census can exceed the assumed 120-bed mean operational capacity during pressure. The source discussion also permits overflow above the regular 100-bed base."),
        ("Should we use 100 or 120 beds in the paper?", "State both clearly: 100 is the regular/base capacity; 120 is the documented mean operational capacity used by MEDCAST because overflow beds can be added. Explain the denominator wherever occupancy is reported."),
        ("Why is Capacity Risk capacity around historical admissions, not 120?", "That control represents daily admission-handling capacity derived from admissions percentiles. It is a flow measure and is separate from bed inventory/census."),
        ("What should we do with a High/Critical risk label?", "Validate the data, selected scenario, prediction interval, occupancy, staff, supplies, and known events; then escalate through authorized hospital planning procedures."),
        ("How often should forecasts be rerun?", "After material new data or corrections, before a formal review if stale, and after analytical changes. Avoid rerunning repeatedly without new information."),
        ("Can MEDCAST predict emergencies or outbreaks?", "It can reflect patterns present in the data but cannot reliably anticipate unprecedented structural shocks without relevant predictors and human context."),
        ("Does interpolation create real observations?", "No. It fills technical gaps for daily modeling. The original source should still be corrected whenever an actual value can be recovered."),
        ("Can the app replace hospital administrators or clinicians?", "No. It organizes evidence and scenarios; responsibility remains with authorized human decision-makers."),
        ("What should be shown in a defense or client presentation?", "Show the dataset version/coverage, 80/20 split, model comparison, one error metric, interval coverage plus width, high-demand performance, sensitivity result, and the 120-bed assumption with its overflow explanation."),
    ]
    for question, answer in faqs:
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(question)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        add_body(doc, answer)

    add_heading(doc, "14. Glossary", 1)
    glossary = [
        ["Actual", "Observed value recorded in the historical dataset"],
        ["AIC / BIC", "Information criteria used to compare SARIMA specifications while penalizing complexity"],
        ["Admissions", "New patient admissions during a day; a flow, not the same as occupied-bed census"],
        ["Best model", "Candidate with the strongest aggregated rolling-origin result under the app’s ranking rule for a horizon"],
        ["Capacity scenario", "Assumed daily admission-handling threshold used for what-if analysis"],
        ["Forecast horizon", "How far ahead the forecast extends: 1, 7, or 30 days"],
        ["Evaluation period", "Newest 20% used as future targets; revealed observations join the expanding history only after their origin date"],
        ["High-demand day", "Actual or forecast admissions above the selected percentile threshold"],
        ["Mean operational capacity", "120-bed planning denominator that incorporates the regular 100-bed base and overflow use"],
        ["Occupied beds", "Number of beds currently occupied by patients"],
        ["Percentile (Pxx)", "Value at or below which approximately xx% of the reference observations fall"],
        ["Point forecast", "Single estimated future value"],
        ["Prediction interval", "Range intended to communicate forecast uncertainty"],
        ["Primary model", "Operational model associated with the central 7-day planning comparison"],
        ["Rolling origin", "Forecast cutoff date that advances through evaluation without using observations beyond that date"],
        ["Robustness", "Stability of the ordered errors produced across rolling origins"],
        ["Initial training set", "Oldest 80% of observations available before the first rolling evaluation origin"],
        ["Winsorization", "Capping extreme values at a percentile for sensitivity analysis"],
    ]
    add_table(doc, ["Term", "Definition"], glossary, [2600, 6760], font_size=8.8)

    add_heading(doc, "15. Client Sign-off Checklist", 1)
    add_body(doc, "Before treating the application and guide as the accepted project baseline, confirm the following with the client:")
    add_bullets(doc, [
        "NDH is the intended hospital scope.",
        "The chronological 80% initial training / 20% expanding-window rolling-origin evaluation is approved.",
        "The regular/base capacity is 100 beds and the mean operational occupancy denominator is 120 beds due to overflow arrangements.",
        "Capacity Risk controls are understood as daily admission-handling capacity, not bed inventory.",
        "The client accepts the Low/Moderate/High percentile definitions and the Capacity Risk score bands.",
        "The client understands that results change when data or assumptions change.",
        "The client accepts the research/decision-support disclaimer and human-approval requirement.",
        "The client has verified the latest source dataset, category definitions, and missing-data handling.",
    ])

    page_break(doc)
    add_heading(doc, "Appendix A. Formula Summary", 1)
    formula_rows = [
        ["Bed occupancy", "Occupied beds ÷ 120 × 100"],
        ["Capacity-pressure ratio", "Average forecast admissions/day ÷ scenario admissions capacity/day"],
        ["Expected overload", "Σ max(point forecast − scenario capacity, 0)"],
        ["Relative interval width", "Average prediction-interval width ÷ mean actual admissions × 100"],
        ["Sensitivity", "TP ÷ (TP + FN) × 100"],
        ["Specificity", "TN ÷ (TN + FP) × 100"],
        ["Precision", "TP ÷ (TP + FP) × 100"],
        ["F1-score", "2 × precision × sensitivity ÷ (precision + sensitivity)"],
        ["False-alert rate", "FP ÷ (FP + TN) × 100"],
        ["Missed-event rate", "FN ÷ (FN + TP) × 100"],
        ["Robustness score", "100 × (1 − min(1, rolling-MAE standard deviation ÷ rolling-MAE mean))"],
        ["Capacity risk score", "[0.65 × exceedance probability + 0.35 × pressure component] × penalty multiplier"],
    ]
    add_table(doc, ["Measure", "Formula / rule"], formula_rows, [3000, 6360], font_size=8.8)

    add_heading(doc, "Appendix B. Current Evaluation Snapshot", 1)
    if snapshot:
        add_body(doc, f"Dataset {snapshot.get('dataset_version', '—')} contains {snapshot.get('dataset_records', '—')} daily observations covering {snapshot.get('dataset_coverage_start', '—')} through {snapshot.get('dataset_coverage_end', '—')}. The evaluation uses {snapshot.get('training_records', '—')} initial training records and {snapshot.get('testing_records', '—')} rolling evaluation days with a {snapshot.get('rolling_origin_step_days', '—')}-day origin step.")
        rows = []
        for item in snapshot.get("benchmarks", []):
            if item.get("is_best_for_horizon"):
                rows.append([
                    str(item.get("horizon_days", "—")),
                    item.get("model_name", "—"),
                    str(item.get("origin_count", "—")),
                    fmt(item.get("mae")),
                    fmt(item.get("rmse")),
                    fmt(item.get("mase"), 3),
                    fmt(item.get("coverage_80"), 1, "%"),
                    fmt(item.get("avg_width_80")),
                    fmt(item.get("robustness_score"), 1, "%"),
                ])
        add_table(doc, ["Days", "Model", "Origins", "MAE", "RMSE", "MASE", "PI80 cov.", "PI80 width", "Robust."], rows, [650, 1350, 800, 950, 950, 900, 1100, 1300, 1360], font_size=7.7)
        evaluation_sarima = snapshot.get("evaluation_sarima_order_selection", {})
        evaluation_selected = evaluation_sarima.get("selected", {}) if isinstance(evaluation_sarima, dict) else {}
        add_body(doc, f"Evaluation SARIMA specification selected from the initial 80%: {evaluation_selected.get('model_order', '—')}. Full-history SARIMA specification for the live forecast: {snapshot.get('selected_sarima_order', '—')}. The web app reports their AIC/BIC metadata separately.", style="Small Note")
    else:
        add_body(doc, "Snapshot unavailable. Rerun the forecasting pipeline and regenerate this guide if a rolling-origin evaluation appendix is required.")

    add_callout(doc, "Final reminder", "The appendix records one local analytical state. The live application remains the authoritative source for the newest dataset version and forecast run.")

    # Apply section header/footer and final pagination preferences.
    for section in doc.sections:
        configure_header_footer(section, first_page=(section is doc.sections[0]))

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build()
    print(path)
